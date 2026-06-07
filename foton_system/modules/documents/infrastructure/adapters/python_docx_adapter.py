import os
import re
from typing import Dict, Any
from docx import Document
from docx.oxml import CT_P
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Cell, _Row
from docx.section import Section
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.document import Document as DocumentType
from lxml.etree import _Element
from foton_system.modules.shared.infrastructure.config.logger import setup_logger
from foton_system.modules.documents.application.ports.document_service_port import DocumentServicePort

logger = setup_logger()


class PythonDocxAdapter(DocumentServicePort):
    """Adapter para manipulação de documentos DOCX via python-docx."""

    def __init__(self) -> None:
        pass

    def load_document(self, path: str) -> DocumentType:
        """Carrega um arquivo DOCX do disco."""
        if not os.path.exists(path):
            raise FileNotFoundError(f"Arquivo DOCX não encontrado: {path}")
        return Document(path)

    def save_document(self, document: DocumentType, path: str) -> None:
        """Persiste o documento DOCX no disco."""
        try:
            document.save(path)
            logger.info(f"Documento salvo em: {path}")
        except Exception as e:
            logger.error(f"Erro ao salvar documento: {e}")
            raise

    def replace_text(self, document: DocumentType, replacements: Dict[str, str]) -> DocumentType:
        """Substitui chaves por valores em parágrafos, tabelas, headers/footers e shapes."""
        logger.info('Iniciando substituição de textos no DOCX...')

        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

        for shape in document.inline_shapes:
            if shape.has_text_frame:
                self._replace_in_text_frame(shape.text_frame, replacements)

        for section in document.sections:
            self._replace_in_header_footer(section.header, replacements)
            self._replace_in_header_footer(section.footer, replacements)

        logger.info('Substituição de textos concluída.')
        return document

    def _replace_in_header_footer(self, header_footer: Any, replacements: Dict[str, str]) -> None:
        """Substitui chaves no header ou footer do documento."""
        if header_footer:
            for paragraph in header_footer.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, replacements)
            self._process_anchored_shapes(header_footer._element, replacements)

    def _replace_in_paragraph(self, paragraph: Paragraph, replacements: Dict[str, str]) -> None:
        """Substitui chaves em um parágrafo, consolidando runs antes."""
        self._consolidate_runs(paragraph)
        for run in paragraph.runs:
            if run.text:
                run.text = self._replace_keys_in_text(run.text, replacements)

    def _consolidate_runs(self, paragraph: Paragraph) -> None:
        """Mescla múltiplos runs para evitar chaves quebradas entre runs."""
        if '@' not in paragraph.text:
            return
        if len(paragraph.runs) > 1:
            full_text: str = "".join(run.text for run in paragraph.runs)
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""

    def _replace_keys_in_text(self, text: str, replacements: Dict[str, str]) -> str:
        """Substitui chaves no texto usando regex case-insensitive."""
        sorted_keys = sorted(replacements.keys(), key=len, reverse=True)
        for key in sorted_keys:
            pattern = r'(?<![\w.])' + re.escape(key) + r'(?!\.[a-z]{2,}\b)'
            new_val: str = str(replacements[key])
            text = re.sub(pattern, new_val, text, flags=re.IGNORECASE)
        return text

    def _replace_in_shapes(self, shapes: Any, replacements: Dict[str, str]) -> None:
        """Substitui chaves em shapes com text_frame."""
        for shape in shapes:
            if shape.has_text_frame:
                self._replace_in_text_frame(shape.text_frame, replacements)

    def _replace_in_text_frame(self, text_frame: Any, replacements: Dict[str, str]) -> None:
        """Substitui chaves em um text_frame."""
        for paragraph in text_frame.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

    def _process_anchored_shapes(self, element: _Element, replacements: Dict[str, str]) -> None:
        """Percorre XML para substituir chaves em text boxes ancoradas."""
        if hasattr(element, "xpath"):
            paragraphs = element.xpath('.//w:txbxContent/w:p')
            for p_element in paragraphs:
                runs = p_element.xpath('.//w:r')
                for run in runs:
                    texts = run.xpath('.//w:t')
                    for t in texts:
                        if t.text:
                            original_text: str = t.text
                            new_text: str = self._replace_keys_in_text(original_text, replacements)
                            if original_text != new_text:
                                t.text = new_text
