import os
from docx import Document
from foton_system.core.logger import setup_logger

logger = setup_logger()

class DOCXHandler:
    def __init__(self):
        pass

    def load_document(self, path):
        if not os.path.exists(path):
            raise FileNotFoundError(f"Arquivo DOCX não encontrado: {path}")
        return Document(path)

    def save_document(self, document, path):
        try:
            document.save(path)
            logger.info(f"Documento salvo em: {path}")
        except Exception as e:
            logger.error(f"Erro ao salvar documento: {e}")
            raise

    def replace_text(self, document, replacements):
        """
        Replaces keys with values in a Word document.
        """
        logger.info('Iniciando substituição de textos no DOCX...')
        
        # Replace in paragraphs
        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        # Replace in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        
        # Replace in headers and footers
        for section in document.sections:
            self._replace_in_header_footer(section.header, replacements)
            self._replace_in_header_footer(section.footer, replacements)

        logger.info('Substituição de textos concluída.')
        return document

    def _replace_in_header_footer(self, header_footer, replacements):
        if header_footer:
            for paragraph in header_footer.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, replacements)

    def _replace_in_paragraph(self, paragraph, replacements):
        # Consolidate runs to fix split keys (e.g. ["@", "key"])
        self._consolidate_runs(paragraph)
        
        for run in paragraph.runs:
            if run.text:
                run.text = self._replace_keys_in_text(run.text, replacements)

    def _consolidate_runs(self, paragraph):
        """
        Consolidates runs in a paragraph to ensure keys aren't split across runs.
        This is a simplified approach: it joins all text and puts it in the first run,
        clearing the others. This preserves text but might lose some formatting 
        if the paragraph has mixed formatting (bold/italic) *within* the text.
        
        However, for templates where keys are usually plain text, this is often acceptable
        and necessary for substitution to work.
        """
        # Check if there's any '@' in the paragraph text first
        if '@' not in paragraph.text:
            return

        # If we have mixed formatting, this is risky. 
        # But if the key itself is split, we have no choice but to join.
        # A safer approach is to only join if we detect a split key, but that's complex.
        # Let's try to join only runs that look like they might be part of a key.
        
        # Actually, for this specific use case (contracts), formatting is usually uniform 
        # within a sentence or the key itself shouldn't have mixed formatting.
        # Let's try the full consolidation for paragraphs containing '@'.
        
        text = paragraph.text
        # Only consolidate if it looks like a key might be present
        if '@' in text:
            # We need to be careful not to destroy formatting of the WHOLE paragraph 
            # if the key is just a small part. 
            # But the 'split run' problem is hard to solve without consolidation.
            
            # Strategy: Join all text, clear all runs, add new run with full text.
            # This resets formatting to the paragraph style, which might be bad.
            # Better: Set text of first run to full text, clear others. 
            # This keeps the formatting of the first run.
            
            # Check if we actually have multiple runs
            if len(paragraph.runs) > 1:
                # Collect all text
                full_text = "".join(run.text for run in paragraph.runs)
                
                # Set first run text
                paragraph.runs[0].text = full_text
                
                # Clear other runs
                for run in paragraph.runs[1:]:
                    run.text = ""
                    # We can't easily remove runs from the list, but setting text to empty works visually.

    def _replace_keys_in_text(self, text, replacements):
        import re
        # Sort keys by length (descending) to avoid partial replacement issues if keys share prefixes
        # e.g. @key and @key2
        sorted_keys = sorted(replacements.keys(), key=len, reverse=True)
        
        for key in sorted_keys:
            if key in text:
                # Use regex to ensure we don't replace inside words/emails
                # Lookbehind: (?<![\w.]) checks if previous char is NOT alphanumeric/underscore/dot
                # Lookahead: (?!\.[a-z]{2,}\b) checks if NOT followed by domain-like suffix
                pattern = r'(?<![\w.])' + re.escape(key) + r'(?!\.[a-z]{2,}\b)'
                
                # We need to replace with the value. 
                # Note: re.sub can take a function or string. String is safer for simple replacement.
                new_val = str(replacements[key])
                text = re.sub(pattern, new_val, text)
                
        return text

    def _replace_in_shapes(self, shapes, replacements):
        for shape in shapes:
            if shape.has_text_frame:
                self._replace_in_text_frame(shape.text_frame, replacements)

    def _replace_in_text_frame(self, text_frame, replacements):
        for paragraph in text_frame.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

    def _process_anchored_shapes(self, element, replacements):
        """
        Traverses XML to find anchored text boxes (w:drawing) and replaces text.
        """
        if hasattr(element, "xpath"):
            # Find all paragraphs within drawings (text boxes)
            # w:txbxContent/w:p is the standard path for text box content
            paragraphs = element.xpath('.//w:txbxContent/w:p')
            for p_element in paragraphs:
                # We need to wrap the XML element in a Paragraph object to use .runs
                # However, creating a Paragraph object from an element is tricky without the parent.
                # A simpler approach for XML manipulation might be needed, 
                # or we can try to find runs directly in XML.
                
                # Let's try to iterate runs (w:r) and text (w:t) directly in XML
                runs = p_element.xpath('.//w:r')
                for run in runs:
                    texts = run.xpath('.//w:t')
                    for t in texts:
                        if t.text:
                            original_text = t.text
                            new_text = self._replace_keys_in_text(original_text, replacements)
                            if original_text != new_text:
                                t.text = new_text

    def replace_text(self, document, replacements):
        """
        Replaces keys with values in a Word document.
        """
        logger.info('Iniciando substituição de textos no DOCX...')
        
        # Replace in paragraphs
        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        # Replace in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        
        # Replace in inline shapes (text boxes in body)
        for shape in document.inline_shapes:
            if shape.has_text_frame:
                self._replace_in_text_frame(shape.text_frame, replacements)

        # Replace in headers and footers (including anchored shapes)
        for section in document.sections:
            self._replace_in_header_footer(section.header, replacements)
            self._replace_in_header_footer(section.footer, replacements)

        logger.info('Substituição de textos concluída.')
        return document

    def _replace_in_header_footer(self, header_footer, replacements):
        if header_footer:
            for paragraph in header_footer.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, replacements)
            
            # Handle anchored shapes in header/footer
            self._process_anchored_shapes(header_footer._element, replacements)
