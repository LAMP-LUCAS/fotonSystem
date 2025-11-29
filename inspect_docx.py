import sys
import os
from docx import Document

# Add project root to path
sys.path.append(os.getcwd())

def inspect_docx():
    file_path = r"c:\Users\Lucas\OneDrive\LAMP_ARQUITETURA\CLIENTES\ED_RIO-AMAZONAS\ACEITE-CBMGO\Proposta_ACEITE-CBMGO.docx"
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    print(f"Inspecting: {file_path}")
    doc = Document(file_path)
    
    print("\n--- Document Content Sample ---")
    count = 0
    for p in doc.paragraphs:
        if p.text.strip():
            print(f"P: {p.text[:100]}...")
            count += 1
            if count > 10: break
            
    print("\n--- Checking for Unreplaced Keys ---")
    keys_found = []
    for p in doc.paragraphs:
        if "@" in p.text:
            keys_found.append(p.text)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "@" in p.text:
                        keys_found.append(p.text)
                        
    if keys_found:
        print(f"Found {len(keys_found)} paragraphs with '@':")
        for k in keys_found[:5]:
            print(f"  - {k.strip()}")
    else:
        print("No keys found (or all replaced).")

if __name__ == "__main__":
    inspect_docx()
